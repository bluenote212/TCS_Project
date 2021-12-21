from docx import Document
from docx.shared import Pt, RGBColor
import aspose.words as aw

doc = aw.Document('C:/Users/B180093/Desktop/release_note.docx')
doc.save("C:/Users/B180093/Desktop/release_note.pdf")



'''
style = document.styles['Normal']
style.font.name = 'Tahoma'

sdk_name = 'None'
sdk_version = 'None'
patch_name = 'None'
patch_version = 'None'
symptom = 'None'
cause = 'None'
solution = 'None'
pre = 'None'
how = 'None'

p = document.paragraphs[1].add_run(sdk_name + '\n' + sdk_version)
p.bold = True
p.font.color.rgb = RGBColor(255, 255, 255)
p.font.size = Pt(20)

for i in document.paragraphs:
    if '{sdk_name}' in i.text:
        i.text = i.text.replace('{sdk_name}', sdk_name)

    elif '{sdk_version}' in i.text:
        i.text = i.text.replace('{sdk_version}', sdk_version)

    elif '{symptom}' in i.text:
        i.text = i.text.replace('{symptom}', 'blablabla')

document.save('C:/Users/B180093/Desktop/result.docx')
'''